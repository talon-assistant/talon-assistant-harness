"""JobTrackerTalent -- track job applications via voice commands.

Stores applications and follow-ups in a dedicated SQLite database
(default: data/job_tracker.db). Supports full CRUD, status tracking,
follow-up reminders, statistics, and XLSX export for unemployment
reporting.

Integration: Cowork can call add_from_cowork() and
get_active_applications() via the cowork_bridge.

Examples:
    "add a job application at Netflix for VP of Engineering"
    "I applied to the Microsoft Azure CISO role"
    "update the Netflix application to interviewing"
    "show my active applications"
    "what jobs need follow up"
    "how many applications this week"
    "mark the Google role as rejected"
    "show all jobs from LinkedIn"
    "export my job tracker"
    "find applications for engineer"
    "add a follow up for Netflix in 5 days"
"""
from __future__ import annotations

import json
import os
import re
import shutil
import sqlite3
import subprocess
import threading
from datetime import datetime, date, timedelta
from pathlib import Path
from typing import Any

from talents.base import BaseTalent
from core.llm_client import LLMError

import logging
log = logging.getLogger(__name__)

# ── Valid statuses and transitions ────────────────────────────────────────────

VALID_STATUSES = ("new", "applied", "interviewing", "offered", "rejected", "withdrawn")

# Each status maps to the set of statuses it can transition TO.
_STATUS_TRANSITIONS: dict[str, set[str]] = {
    "new":          {"applied", "withdrawn"},
    "applied":      {"interviewing", "offered", "rejected", "withdrawn"},
    "interviewing": {"offered", "rejected", "withdrawn"},
    "offered":      {"rejected", "withdrawn"},
    "rejected":     set(),
    "withdrawn":    set(),
}


def _data_dir() -> str:
    """Ensure data/ directory exists and return its path."""
    d = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data")
    os.makedirs(d, exist_ok=True)
    return d


def _normalize_company(name: str) -> str:
    """Normalize a company name for fuzzy matching."""
    n = name.lower().strip()
    # Strip common suffixes
    for suffix in (" inc", " inc.", " llc", " corp", " corp.",
                   " co", " co.", " ltd", " ltd.", " limited",
                   " corporation", " incorporated"):
        if n.endswith(suffix):
            n = n[: -len(suffix)].rstrip(",. ")
    # Strip leading "the"
    if n.startswith("the "):
        n = n[4:]
    return n.strip()


def _friendly_date(iso_str: str | None) -> str:
    """Format an ISO date string as a human-friendly string."""
    if not iso_str:
        return ""
    try:
        d = date.fromisoformat(iso_str[:10])
    except (ValueError, TypeError):
        return iso_str
    today = date.today()
    delta = (today - d).days
    if delta == 0:
        return "today"
    elif delta == 1:
        return "yesterday"
    elif delta < 7:
        return f"{delta} days ago"
    else:
        return d.strftime("%b %d")


def _today_iso() -> str:
    return date.today().isoformat()


# ── Database helpers ──────────────────────────────────────────────────────────

_SCHEMA_APPLICATIONS = """\
CREATE TABLE IF NOT EXISTS applications (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    company TEXT NOT NULL,
    position TEXT NOT NULL,
    location TEXT DEFAULT '',
    source TEXT DEFAULT '',
    status TEXT DEFAULT 'new',
    date_found TEXT,
    date_applied TEXT,
    date_updated TEXT,
    contact_name TEXT DEFAULT '',
    contact_email TEXT DEFAULT '',
    method TEXT DEFAULT '',
    salary_range TEXT DEFAULT '',
    notes TEXT DEFAULT '',
    job_url TEXT DEFAULT '',
    resume_version TEXT DEFAULT '',
    cover_letter INTEGER DEFAULT 0,
    cowork_task_id TEXT DEFAULT '',
    fit_score INTEGER DEFAULT 0,
    archived INTEGER DEFAULT 0
)"""

_SCHEMA_FOLLOW_UPS = """\
CREATE TABLE IF NOT EXISTS follow_ups (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    application_id INTEGER NOT NULL,
    due_date TEXT NOT NULL,
    note TEXT DEFAULT '',
    completed INTEGER DEFAULT 0,
    FOREIGN KEY (application_id) REFERENCES applications(id)
)"""


class _DB:
    """Thin SQLite wrapper scoped to a single database file."""

    def __init__(self, db_path: str) -> None:
        self.db_path = db_path
        os.makedirs(os.path.dirname(db_path) or ".", exist_ok=True)
        self._init_schema()

    def _connect(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA foreign_keys = ON")
        return conn

    def _init_schema(self) -> None:
        with self._connect() as conn:
            conn.execute(_SCHEMA_APPLICATIONS)
            conn.execute(_SCHEMA_FOLLOW_UPS)
            # Idempotent migrations for columns added after initial release.
            cols = {r[1] for r in conn.execute(
                "PRAGMA table_info(applications)"
            ).fetchall()}
            if "job_description" not in cols:
                conn.execute(
                    "ALTER TABLE applications "
                    "ADD COLUMN job_description TEXT DEFAULT ''"
                )
            if "recruiter_name" not in cols:
                conn.execute(
                    "ALTER TABLE applications "
                    "ADD COLUMN recruiter_name TEXT DEFAULT ''"
                )
            if "recruiter_url" not in cols:
                conn.execute(
                    "ALTER TABLE applications "
                    "ADD COLUMN recruiter_url TEXT DEFAULT ''"
                )
            if "connections_at_co" not in cols:
                conn.execute(
                    "ALTER TABLE applications "
                    "ADD COLUMN connections_at_co TEXT DEFAULT ''"
                )
            if "archived_at" not in cols:
                conn.execute(
                    "ALTER TABLE applications "
                    "ADD COLUMN archived_at TEXT"
                )
                # Grace backfill: any rows that were already archived
                # before this column existed get stamped NOW so they
                # have a fresh 30-day clock instead of being instantly
                # eligible for expiry.
                conn.execute(
                    "UPDATE applications "
                    "SET archived_at = datetime('now') "
                    "WHERE archived = 1 AND archived_at IS NULL"
                )

    # -- applications --

    def add_application(self, **kwargs: Any) -> int:
        """Insert a new application row. Returns the new row id."""
        import html as _html
        for _f in ("company", "position", "location"):
            if kwargs.get(_f):
                kwargs[_f] = _html.unescape(str(kwargs[_f])).strip()
        cols = [k for k in kwargs if kwargs[k] is not None]
        placeholders = ", ".join("?" for _ in cols)
        col_names = ", ".join(cols)
        vals = [kwargs[c] for c in cols]
        with self._connect() as conn:
            cur = conn.execute(
                f"INSERT INTO applications ({col_names}) VALUES ({placeholders})", vals
            )
            return cur.lastrowid  # type: ignore[return-value]

    def unescape_html_entities(self) -> int:
        """Idempotent backfill: unescape HTML entities in text fields.

        Returns the number of rows modified. Safe to run repeatedly.
        """
        import html as _html
        fields = ("company", "position", "location")
        changed = 0
        with self._connect() as conn:
            rows = conn.execute(
                f"SELECT id, {', '.join(fields)} FROM applications"
            ).fetchall()
            for row in rows:
                updates: dict[str, str] = {}
                for f in fields:
                    val = row[f]
                    if not val:
                        continue
                    clean = _html.unescape(str(val)).strip()
                    if clean != val:
                        updates[f] = clean
                if updates:
                    sets = ", ".join(f"{k} = ?" for k in updates)
                    vals = list(updates.values()) + [row["id"]]
                    conn.execute(
                        f"UPDATE applications SET {sets} WHERE id = ?", vals
                    )
                    changed += 1
        return changed

    def update_application(self, app_id: int, **kwargs: Any) -> bool:
        """Update fields on an existing application. Returns True if a row was modified."""
        sets = ", ".join(f"{k} = ?" for k in kwargs)
        vals = list(kwargs.values()) + [app_id]
        with self._connect() as conn:
            cur = conn.execute(
                f"UPDATE applications SET {sets} WHERE id = ? AND archived = 0", vals
            )
            return cur.rowcount > 0

    def get_application(self, app_id: int) -> dict | None:
        with self._connect() as conn:
            row = conn.execute(
                "SELECT * FROM applications WHERE id = ? AND archived = 0", (app_id,)
            ).fetchone()
            return dict(row) if row else None

    def find_by_company(self, company: str) -> list[dict]:
        """Fuzzy-find applications by company name (non-archived)."""
        norm = _normalize_company(company)
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM applications WHERE archived = 0 ORDER BY id DESC"
            ).fetchall()
        return [dict(r) for r in rows if _normalize_company(r["company"]) == norm]

    def search(self, term: str) -> list[dict]:
        """Search company and position columns."""
        pattern = f"%{term}%"
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM applications WHERE archived = 0 "
                "AND (company LIKE ? OR position LIKE ?) "
                "ORDER BY id DESC",
                (pattern, pattern),
            ).fetchall()
        return [dict(r) for r in rows]

    def list_active(self) -> list[dict]:
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM applications WHERE archived = 0 "
                "AND status NOT IN ('rejected', 'withdrawn') "
                "ORDER BY date_updated DESC, id DESC"
            ).fetchall()
        return [dict(r) for r in rows]

    def list_by_status(self, status: str) -> list[dict]:
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM applications WHERE archived = 0 AND status = ? "
                "ORDER BY date_updated DESC", (status,)
            ).fetchall()
        return [dict(r) for r in rows]

    def list_by_source(self, source: str) -> list[dict]:
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM applications WHERE archived = 0 "
                "AND LOWER(source) = LOWER(?) ORDER BY id DESC", (source,)
            ).fetchall()
        return [dict(r) for r in rows]

    def list_top_candidates(self, limit: int = 15) -> list[dict]:
        """Return new applications sorted by fit_score descending."""
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM applications WHERE archived = 0 "
                "AND status = 'new' AND fit_score > 0 "
                "ORDER BY fit_score DESC, id DESC LIMIT ?",
                (limit,),
            ).fetchall()
        return [dict(r) for r in rows]

    def list_all(self, include_archived: bool = False) -> list[dict]:
        clause = "" if include_archived else "WHERE archived = 0"
        with self._connect() as conn:
            rows = conn.execute(
                f"SELECT * FROM applications {clause} ORDER BY id DESC"
            ).fetchall()
        return [dict(r) for r in rows]

    def count_since(self, since_iso: str) -> int:
        with self._connect() as conn:
            row = conn.execute(
                "SELECT COUNT(*) FROM applications WHERE archived = 0 "
                "AND date_found >= ?", (since_iso,)
            ).fetchone()
            return row[0] if row else 0

    def stats(self) -> dict[str, int]:
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT status, COUNT(*) FROM applications WHERE archived = 0 "
                "GROUP BY status"
            ).fetchall()
        return {r[0]: r[1] for r in rows}

    def archive(self, app_id: int) -> bool:
        with self._connect() as conn:
            cur = conn.execute(
                "UPDATE applications "
                "SET archived = 1, archived_at = datetime('now') "
                "WHERE id = ?",
                (app_id,)
            )
            return cur.rowcount > 0

    def unarchive(self, app_id: int) -> bool:
        """Restore an archived application (clears archived_at)."""
        with self._connect() as conn:
            cur = conn.execute(
                "UPDATE applications "
                "SET archived = 0, archived_at = NULL "
                "WHERE id = ?",
                (app_id,)
            )
            return cur.rowcount > 0

    def expire_old_archives(
        self, days: int = 30,
        audit_log_path: str | None = None,
    ) -> dict:
        """Hard-delete archived applications older than `days`.

        Args:
            days:           Age threshold. Rows with archived_at older
                            than this are removed.
            audit_log_path: If set, append a CSV row per deleted entry
                            so the user can recover company/position
                            info if they ever wonder what was purged.

        Returns:
            {expired: int, audit_path: str | None}
        """
        if days <= 0:
            return {"expired": 0, "audit_path": None}

        with self._connect() as conn:
            # Fetch the rows we're about to delete so we can audit-log
            # and clean up follow-ups in the same transaction.
            rows = conn.execute(
                "SELECT id, company, position, source, date_found, "
                "       date_applied, archived_at "
                "FROM applications "
                "WHERE archived = 1 "
                "  AND archived_at IS NOT NULL "
                "  AND archived_at < datetime('now', ?)",
                (f"-{int(days)} days",),
            ).fetchall()
            if not rows:
                return {"expired": 0, "audit_path": None}

            ids = [r["id"] for r in rows]
            placeholders = ",".join("?" * len(ids))
            conn.execute(
                f"DELETE FROM follow_ups WHERE application_id IN ({placeholders})",
                ids,
            )
            conn.execute(
                f"DELETE FROM applications WHERE id IN ({placeholders})",
                ids,
            )

        # Audit log written outside the DB transaction so a write failure
        # doesn't roll back the cleanup itself.
        if audit_log_path and rows:
            try:
                import csv
                from datetime import datetime as _dt
                exists = os.path.exists(audit_log_path)
                with open(audit_log_path, "a", newline="",
                          encoding="utf-8") as f:
                    w = csv.writer(f)
                    if not exists:
                        w.writerow([
                            "expired_at", "id", "company", "position",
                            "source", "date_found", "date_applied",
                            "archived_at",
                        ])
                    now_iso = _dt.now().isoformat(timespec="seconds")
                    for r in rows:
                        w.writerow([
                            now_iso, r["id"], r["company"], r["position"],
                            r["source"] or "", r["date_found"] or "",
                            r["date_applied"] or "", r["archived_at"],
                        ])
            except Exception as exc:
                log.warning(
                    f"[JobTracker] Audit log write failed "
                    f"({audit_log_path}): {exc}"
                )

        return {"expired": len(rows), "audit_path": audit_log_path}

    def hard_delete(self, app_id: int) -> bool:
        """Permanently delete an application row and its follow-ups."""
        with self._connect() as conn:
            conn.execute(
                "DELETE FROM follow_ups WHERE application_id = ?", (app_id,)
            )
            cur = conn.execute(
                "DELETE FROM applications WHERE id = ?", (app_id,)
            )
            return cur.rowcount > 0

    # -- follow-ups --

    def add_follow_up(self, application_id: int, due_date: str,
                      note: str = "") -> int:
        with self._connect() as conn:
            cur = conn.execute(
                "INSERT INTO follow_ups (application_id, due_date, note) "
                "VALUES (?, ?, ?)",
                (application_id, due_date, note),
            )
            return cur.lastrowid  # type: ignore[return-value]

    def get_pending_follow_ups(self) -> list[dict]:
        """Return incomplete follow-ups with application info, ordered by due date."""
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT f.*, a.company, a.position FROM follow_ups f "
                "JOIN applications a ON f.application_id = a.id "
                "WHERE f.completed = 0 AND a.archived = 0 "
                "ORDER BY f.due_date ASC"
            ).fetchall()
        return [dict(r) for r in rows]

    def get_overdue_follow_ups(self) -> list[dict]:
        today = _today_iso()
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT f.*, a.company, a.position FROM follow_ups f "
                "JOIN applications a ON f.application_id = a.id "
                "WHERE f.completed = 0 AND f.due_date <= ? AND a.archived = 0 "
                "ORDER BY f.due_date ASC",
                (today,),
            ).fetchall()
        return [dict(r) for r in rows]

    def complete_follow_up(self, follow_up_id: int) -> bool:
        with self._connect() as conn:
            cur = conn.execute(
                "UPDATE follow_ups SET completed = 1 WHERE id = ?", (follow_up_id,)
            )
            return cur.rowcount > 0


# ── LLM extraction prompt ────────────────────────────────────────────────────

_EXTRACT_SYSTEM = """\
You are a structured data extractor for a job application tracker.
Given a user command, extract the requested fields as a JSON object.
Return ONLY valid JSON, no markdown, no explanation.
If a field cannot be determined, omit it from the JSON."""

_EXTRACT_ADD_PROMPT = """\
Extract job application details from this command.

Fields to extract:
- company (string, REQUIRED)
- position (string, REQUIRED)
- source (string: linkedin, indeed, recruiter, direct, referral, etc.)
- location (string)
- method (string: online, email, recruiter, referral)
- salary_range (string)
- already_applied (boolean: true if user says they already applied)

Command: {command}"""

_EXTRACT_UPDATE_PROMPT = """\
Extract the company name and new status from this command.

Valid statuses: new, applied, interviewing, offered, rejected, withdrawn

Fields:
- company (string, REQUIRED)
- status (string from list above, REQUIRED)

Command: {command}"""

_EXTRACT_FOLLOW_UP_PROMPT = """\
Extract follow-up details from this command.

Fields:
- company (string, REQUIRED)
- days (integer: number of days from now for the follow-up, default 3)
- note (string: what the follow-up is about)

Command: {command}"""


# ── Talent ────────────────────────────────────────────────────────────────────

class JobTrackerTalent(BaseTalent):
    """Track job applications, follow-ups, and export reports."""

    name = "job_tracker"
    description = (
        "Track job applications, update statuses, manage follow-ups, "
        "view statistics, and export reports for tracked applications"
    )
    keywords = [
        "application", "applied", "job tracker", "tracker",
        "follow up", "interview", "applications",
        "active applications", "job application",
        "export my job", "import my job",
        "cover letter", "top jobs", "best matches",
        "top candidates", "best jobs",
    ]
    examples = [
        "add a job application at Netflix for VP of Engineering",
        "I applied to the Microsoft Azure CISO role",
        "update the Netflix application to interviewing",
        "show my active applications",
        "show top jobs",
        "show best matches",
        "write a cover letter for the Affirm job",
        "what jobs need follow up",
        "how many applications this week",
        "export my job tracker",
    ]
    priority = 55

    def __init__(self) -> None:
        super().__init__()
        self._db: _DB | None = None

    def initialize(self, config: dict) -> None:
        """Open (or create) the job tracker database."""
        db_path = self.talent_config.get(
            "db_path",
            os.path.join(_data_dir(), "job_tracker.db"),
        )
        try:
            self._db = _DB(db_path)
            log.info(f"[JobTracker] Database ready at {db_path}")
        except Exception as e:
            log.error(f"[JobTracker] Failed to open database: {e}")
            self._db = None
            return

        # Auto-expire archived applications older than the configured
        # retention window. Keeps the inbox table small so the UI stays
        # snappy at scale. 0 disables; default 30 days.
        retention_days = int(self.talent_config.get(
            "archive_retention_days", 30))
        if retention_days > 0:
            audit_path = os.path.join(
                os.path.dirname(db_path), "job_archive_purge.csv")
            try:
                result = self._db.expire_old_archives(
                    days=retention_days, audit_log_path=audit_path)
                if result.get("expired", 0) > 0:
                    log.info(
                        f"[JobTracker] Auto-expired {result['expired']} "
                        f"archived application(s) older than "
                        f"{retention_days} days. Audit: {audit_path}"
                    )
            except Exception as exc:
                log.warning(f"[JobTracker] Archive cleanup failed: {exc}")

    @property
    def routing_available(self) -> bool:
        return self._db is not None

    # ── Config schema ─────────────────────────────────────────────────────────

    def get_config_schema(self) -> dict:
        return {
            "fields": [
                {
                    "key": "db_path",
                    "label": "Database Path",
                    "type": "string",
                    "default": os.path.join(_data_dir(), "job_tracker.db"),
                },
                {
                    "key": "archive_retention_days",
                    "label": "Archive Retention (days)",
                    "type": "int",
                    "default": 30,
                    "help": (
                        "Auto-purge archived applications older than this "
                        "many days on Talon startup. 0 disables. Removed "
                        "rows are appended to job_archive_purge.csv next "
                        "to the database file for audit/recovery."
                    ),
                },
            ]
        }

    # ── Routing ───────────────────────────────────────────────────────────────

    def can_handle(self, command: str) -> bool:
        return self.keyword_match(command)

    # ── Main dispatch ─────────────────────────────────────────────────────────

    def execute(self, command: str, context: dict) -> dict:
        """Route the user command to the appropriate handler."""
        cmd = command.lower()

        # Determine intent from keywords in the command
        if "import" in cmd and ("spreadsheet" in cmd or "xlsx" in cmd or "excel" in cmd):
            return self._handle_import(command, context)
        if self._is_cover_letter(cmd):
            return self._handle_cover_letter(command, context)
        if self._is_top_candidates(cmd):
            return self._handle_top_candidates()
        if self._is_export(cmd):
            return self._handle_export(context)
        if self._is_follow_up(cmd):
            if self._is_query(cmd):
                return self._handle_follow_up_list()
            return self._handle_follow_up_add(command, context)
        if self._is_stats(cmd):
            return self._handle_stats(cmd)
        if self._is_update(cmd):
            return self._handle_update(command, context)
        # "#ID applied" / "applied for #7" — update, not add
        if re.search(r'#\d+', cmd) and any(s in cmd for s in VALID_STATUSES):
            return self._handle_update(command, context)
        if self._is_add(cmd):
            return self._handle_add(command, context)
        if self._is_search(cmd):
            return self._handle_search(command, context)
        if self._is_list(cmd):
            return self._handle_list(cmd)

        # Fallback: try to figure out intent via LLM
        intent = self._extract_arg(
            context["llm"], command, "intent",
            options=["add", "update", "list", "stats", "export",
                     "follow_up", "search", "cover_letter", "top_candidates"],
        )
        handlers = {
            "add": lambda: self._handle_add(command, context),
            "update": lambda: self._handle_update(command, context),
            "list": lambda: self._handle_list(cmd),
            "stats": lambda: self._handle_stats(cmd),
            "export": lambda: self._handle_export(context),
            "follow_up": lambda: self._handle_follow_up_add(command, context),
            "search": lambda: self._handle_search(command, context),
            "cover_letter": lambda: self._handle_cover_letter(command, context),
            "top_candidates": lambda: self._handle_top_candidates(),
        }
        if intent and intent in handlers:
            return handlers[intent]()

        # Ultimate fallback: show active applications
        return self._handle_list(cmd)

    # ── Intent detection helpers ──────────────────────────────────────────────

    @staticmethod
    def _is_add(cmd: str) -> bool:
        return bool(re.search(
            r'\b(add|new|applied to|i applied|applied for|submit)\b', cmd
        ))

    @staticmethod
    def _is_update(cmd: str) -> bool:
        return bool(re.search(
            r'\b(update|mark|change status|move to|set status|got an offer|got rejected)\b',
            cmd,
        ))

    @staticmethod
    def _is_list(cmd: str) -> bool:
        return bool(re.search(
            r'\b(show|list|active|display|view|what are|my applications)\b', cmd
        ))

    @staticmethod
    def _is_stats(cmd: str) -> bool:
        return bool(re.search(
            r'\b(how many|count|stats|statistics|summary|this week|this month)\b', cmd
        ))

    @staticmethod
    def _is_export(cmd: str) -> bool:
        return bool(re.search(
            r'\b(export|spreadsheet|xlsx|generate report|download)\b', cmd
        ))

    @staticmethod
    def _is_follow_up(cmd: str) -> bool:
        return bool(re.search(r'\bfollow[ -]?up\b', cmd))

    @staticmethod
    def _is_query(cmd: str) -> bool:
        return bool(re.search(
            r'\b(show|list|what|need|overdue|pending|check)\b', cmd
        ))

    @staticmethod
    def _is_search(cmd: str) -> bool:
        return bool(re.search(r'\b(find|search|look for|lookup)\b', cmd))

    @staticmethod
    def _is_top_candidates(cmd: str) -> bool:
        return bool(re.search(
            r'\b(top jobs|best match|best jobs|top candidates|top matches|'
            r'highest fit|best fit|ranked|strongest)\b', cmd
        ))

    @staticmethod
    def _is_cover_letter(cmd: str) -> bool:
        return bool(re.search(r'\bcover\s*letter\b', cmd))

    # ── Add ───────────────────────────────────────────────────────────────────

    def _handle_add(self, command: str, context: dict) -> dict:
        """Add a new job application."""
        llm = context["llm"]
        data = self._extract_json(llm, _EXTRACT_ADD_PROMPT.format(command=command))
        if not data or not data.get("company") or not data.get("position"):
            return self._fail("I need at least a company and position. "
                              "Try: 'add a job at Netflix for Senior Engineer'")

        today = _today_iso()
        already_applied = data.get("already_applied", False)

        row = {
            "company": data["company"],
            "position": data["position"],
            "location": data.get("location", ""),
            "source": data.get("source", ""),
            "method": data.get("method", ""),
            "salary_range": data.get("salary_range", ""),
            "status": "applied" if already_applied else "new",
            "date_found": today,
            "date_applied": today if already_applied else "",
            "date_updated": today,
        }

        app_id = self._db.add_application(**row)
        status_note = " (marked as applied)" if already_applied else ""

        return {
            "success": True,
            "response": (
                f"Added application #{app_id}: **{data['position']}** at "
                f"**{data['company']}**{status_note}."
            ),
            "actions_taken": [{"action": "job_add", "id": app_id,
                               "company": data["company"]}],
            "spoken": False,
        }

    # ── Update status ─────────────────────────────────────────────────────────

    def _handle_update(self, command: str, context: dict) -> dict:
        """Update the status of an existing application."""
        cmd_lower = command.lower()

        # Fast path: "#ID applied" / "applied for #ID" / "mark #ID as applied"
        id_match = re.search(r'#(\d+)', command)
        if id_match:
            app_id = int(id_match.group(1))
            app = self._db.get_application(app_id)
            if app:
                # Figure out the target status from the command text
                new_status = None
                for s in VALID_STATUSES:
                    if s in cmd_lower:
                        new_status = s
                        break
                if not new_status:
                    # Default: if "applied" is anywhere in the command
                    if "appli" in cmd_lower:
                        new_status = "applied"
                if new_status:
                    return self._apply_status_update(app, new_status)
                return self._fail(
                    f"What status? Try: 'mark #{app_id} as applied'"
                )

        llm = context["llm"]
        data = self._extract_json(llm, _EXTRACT_UPDATE_PROMPT.format(command=command))
        if not data or not data.get("company") or not data.get("status"):
            return self._fail("I need a company name and new status. "
                              "Try: 'mark #7 as applied' or "
                              "'update Sony to applied'")

        new_status = data["status"].lower().strip()
        if new_status not in VALID_STATUSES:
            return self._fail(
                f"'{new_status}' is not a valid status. "
                f"Options: {', '.join(VALID_STATUSES)}"
            )

        matches = self._db.find_by_company(data["company"])
        if not matches:
            # Try a broader search
            matches = self._db.search(data["company"])
        if not matches:
            return self._fail(
                f"No application found matching '{data['company']}'. "
                "Use 'show my applications' to see what's tracked."
            )

        app = matches[0]  # Most recent match
        return self._apply_status_update(app, new_status)

    def _apply_status_update(self, app: dict, new_status: str) -> dict:
        """Apply a status change to an application with transition validation."""
        current = app["status"]
        allowed = _STATUS_TRANSITIONS.get(current, set())
        if new_status != current and new_status not in allowed:
            return self._fail(
                f"Cannot move from '{current}' to '{new_status}'. "
                f"Allowed transitions from '{current}': "
                f"{', '.join(sorted(allowed)) or 'none (terminal status)'}."
            )

        update_fields: dict[str, Any] = {
            "status": new_status,
            "date_updated": _today_iso(),
        }
        if new_status == "applied" and not app.get("date_applied"):
            update_fields["date_applied"] = _today_iso()

        self._db.update_application(app["id"], **update_fields)

        # When moving to "applied" on a LinkedIn job, fire background
        # "I'm Interested" click on the company's LinkedIn page.
        if (new_status == "applied"
                and app.get("source", "").lower() == "linkedin"
                and app.get("job_url")):
            import threading
            threading.Thread(
                target=self._linkedin_im_interested,
                args=(app["job_url"], app["company"]),
                daemon=True,
                name="linkedin-interested",
            ).start()
            log.info(f"[JobTracker] LinkedIn 'I'm Interested' click queued "
                     f"for {app['company']}")

        return {
            "success": True,
            "response": (
                f"Updated **{app['company']}** ({app['position']}) "
                f"from *{current}* to *{new_status}*."
            ),
            "actions_taken": [{"action": "job_update", "id": app["id"],
                               "from": current, "to": new_status}],
            "spoken": False,
        }

    # ── List ──────────────────────────────────────────────────────────────────

    def _handle_list(self, cmd: str) -> dict:
        """List applications, with optional filtering."""
        # Check for source filter: "from linkedin"
        source_match = re.search(r'\bfrom\s+(\w+)', cmd)
        if source_match:
            source = source_match.group(1)
            apps = self._db.list_by_source(source)
            title = f"Applications from {source}"
        # Check for status filter
        elif any(s in cmd for s in VALID_STATUSES):
            status = next(s for s in VALID_STATUSES if s in cmd)
            apps = self._db.list_by_status(status)
            title = f"Applications with status '{status}'"
        else:
            apps = self._db.list_active()
            title = "Active applications"

        if not apps:
            return {
                "success": True,
                "response": f"No applications found ({title.lower()}).",
                "actions_taken": [{"action": "job_list"}],
                "spoken": False,
            }

        lines = [f"**{title}** ({len(apps)}):\n"]
        for app in apps:
            status_icon = {
                "new": "o", "applied": ">", "interviewing": "?",
                "offered": "$", "rejected": "x", "withdrawn": "-",
            }.get(app["status"], " ")
            date_str = _friendly_date(app.get("date_updated") or app.get("date_found"))
            lines.append(
                f"[{status_icon}] #{app['id']} **{app['company']}** -- "
                f"{app['position']} ({app['status']}) {date_str}"
            )

        return {
            "success": True,
            "response": "\n".join(lines),
            "actions_taken": [{"action": "job_list", "count": len(apps)}],
            "spoken": False,
        }

    # ── Stats ─────────────────────────────────────────────────────────────────

    def _handle_stats(self, cmd: str) -> dict:
        """Show application statistics."""
        status_counts = self._db.stats()
        total = sum(status_counts.values())

        today = date.today()
        week_start = (today - timedelta(days=today.weekday())).isoformat()
        month_start = today.replace(day=1).isoformat()

        this_week = self._db.count_since(week_start)
        this_month = self._db.count_since(month_start)

        lines = [f"**Job Search Stats** (total: {total})\n"]
        for status in VALID_STATUSES:
            count = status_counts.get(status, 0)
            if count > 0:
                lines.append(f"  {status}: {count}")
        lines.append(f"\nThis week: {this_week}")
        lines.append(f"This month: {this_month}")

        # Overdue follow-ups
        overdue = self._db.get_overdue_follow_ups()
        if overdue:
            lines.append(f"\n{len(overdue)} overdue follow-up(s)!")

        return {
            "success": True,
            "response": "\n".join(lines),
            "actions_taken": [{"action": "job_stats"}],
            "spoken": False,
        }

    # ── Search ────────────────────────────────────────────────────────────────

    def _handle_search(self, command: str, context: dict) -> dict:
        """Search applications by company or position text."""
        llm = context["llm"]
        term = self._extract_arg(llm, command, "search term") or ""
        if not term:
            # Try stripping common prefixes
            for prefix in ("find", "search", "look for", "lookup",
                           "find applications for", "search for"):
                if command.lower().startswith(prefix):
                    term = command[len(prefix):].strip()
                    break
        if not term:
            return self._fail("What should I search for? "
                              "Try: 'find applications for engineer'")

        apps = self._db.search(term)
        if not apps:
            return {
                "success": True,
                "response": f"No applications matching '{term}'.",
                "actions_taken": [{"action": "job_search", "term": term}],
                "spoken": False,
            }

        lines = [f"**Search results for '{term}'** ({len(apps)}):\n"]
        for app in apps:
            lines.append(
                f"  #{app['id']} **{app['company']}** -- {app['position']} "
                f"({app['status']})"
            )

        return {
            "success": True,
            "response": "\n".join(lines),
            "actions_taken": [{"action": "job_search", "term": term,
                               "count": len(apps)}],
            "spoken": False,
        }

    # ── Top candidates ───────────────────────────────────────────────────────

    def _handle_top_candidates(self) -> dict:
        """Show new jobs ranked by fit score."""
        apps = self._db.list_top_candidates(limit=15)
        if not apps:
            return {
                "success": True,
                "response": (
                    "No scored candidates yet. Fit scores are added "
                    "automatically after a job search runs. Try 'search for "
                    "jobs' first, then check back in a few minutes."
                ),
                "actions_taken": [{"action": "top_candidates"}],
                "spoken": False,
            }

        lines = [f"**Top candidates** ({len(apps)}):\n"]
        for app in apps:
            loc = f" ({app['location']})" if app.get("location") else ""
            rec = ""
            notes = app.get("notes", "")
            if "Recommendation:" in notes:
                rec_part = notes.split("Recommendation:")[-1].strip()
                rec = f" [{rec_part}]"
            url = app.get("job_url", "")
            url_part = f" | {url}" if url else ""
            lines.append(
                f"  #{app['id']} [{app['fit_score']}%] "
                f"**{app['company']}** -- {app['position']}{loc}{rec}"
                f"{url_part}"
            )

        lines.append(
            "\nSay 'write a cover letter for #ID' or "
            "'write a cover letter for [company]' to generate one."
        )

        return {
            "success": True,
            "response": "\n".join(lines),
            "actions_taken": [{"action": "top_candidates", "count": len(apps)}],
            "spoken": False,
        }

    # ── Cover letter generation ──────────────────────────────────────────────

    def _handle_cover_letter(self, command: str, context: dict) -> dict:
        """Generate a cover letter for a specific application via Claude CLI."""
        # Find the application - try ID first, then company name
        app = None
        id_match = re.search(r'#?(\d+)', command)
        if id_match:
            app = self._db.get_application(int(id_match.group(1)))

        if not app:
            # Extract company name via LLM
            llm = context["llm"]
            company = self._extract_arg(
                llm, command, "company name for the cover letter"
            )
            if company:
                matches = self._db.find_by_company(company)
                if not matches:
                    matches = self._db.search(company)
                if matches:
                    app = matches[0]

        if not app:
            return self._fail(
                "Which job? Try: 'write a cover letter for #12' "
                "or 'write a cover letter for Affirm'"
            )

        # Read resume inline
        resume_path = Path.home() / "OneDrive" / "Documents" / "resume_bullet_library.md"
        try:
            resume_text = resume_path.read_text(encoding="utf-8")
        except Exception as e:
            return self._fail(f"Cannot read resume: {e}")

        # Read CLAUDE.md for style rules
        claude_md_path = Path.home() / ".claude" / "CLAUDE.md"
        style_rules = ""
        try:
            style_rules = claude_md_path.read_text(encoding="utf-8")
        except Exception:
            pass

        # Scrape job description if URL available
        job_url = app.get("job_url", "")
        job_description = ""
        if job_url:
            job_description = self._fetch_job_description(job_url)

        prompt_parts = [
            "TASK: Write a cover letter for the position below.",
            "",
            f"RESUME:\n{resume_text}",
            "",
            f"POSITION: {app['position']}",
            f"COMPANY: {app['company']}",
        ]
        if app.get("location"):
            prompt_parts.append(f"LOCATION: {app['location']}")
        if job_description:
            prompt_parts.append(f"\nJOB DESCRIPTION:\n{job_description}")
        if app.get("notes"):
            fit_notes = app["notes"]
            if "Recommendation:" in fit_notes:
                prompt_parts.append(f"\nFIT ANALYSIS: {fit_notes}")

        prompt_parts.extend([
            "",
            f"STYLE RULES:\n{style_rules}" if style_rules else "",
            "",
            "INSTRUCTIONS:",
            "- Only promote matches and strengths. NEVER mention "
            "weaknesses, gaps, or missing qualifications.",
            "- No em dashes. Use commas, periods, or semicolons.",
            "- No tricolon / parallel triplets.",
            "- Plain language, direct executive tone.",
            "- Open with a specific hook tied to the company or role.",
            "- Pull specific metrics and accomplishments from the resume "
            "that match this role.",
            "- 3-4 paragraphs, under one page.",
            "- Close with confidence, not desperation.",
            "",
            "Output ONLY the cover letter text. No commentary, no "
            "preamble like 'Here is your cover letter', no markdown. "
            "Just the letter itself, starting with 'Dear'.",
        ])

        prompt = "\n".join(p for p in prompt_parts if p is not None)

        # Run Claude CLI
        claude_bin = shutil.which("claude")
        if not claude_bin:
            return self._fail(
                "Claude CLI not found. Install with: "
                "npm i -g @anthropic-ai/claude-code"
            )
        try:
            result = subprocess.run(
                [claude_bin, "-p", "--output-format", "text"],
                input=prompt,
                capture_output=True,
                text=True,
                encoding="utf-8",
                timeout=120,
                cwd=str(Path.home()),
            )

            if result.returncode != 0:
                log.error(
                    f"[JobTracker] claude -p cover letter failed: "
                    f"{result.stderr[:200]}"
                )
                return self._fail(
                    "Claude CLI failed to generate the cover letter. "
                    "Check that 'claude' is installed and working."
                )

            letter = result.stdout.strip()
            if not letter or len(letter) < 50:
                return self._fail("Claude returned an empty or too-short response.")

            # Sanity check — reject if Claude returned an error instead
            # of a cover letter
            _bad_signals = ["webfetch", "permission", "tool call",
                            "blocked by", "approve the"]
            if any(s in letter.lower() for s in _bad_signals):
                log.error(f"[JobTracker] Cover letter looks like an error: "
                          f"{letter[:200]}")
                return self._fail(
                    "Claude returned an error instead of a cover letter. "
                    "This has been fixed — try again."
                )

        except subprocess.TimeoutExpired:
            return self._fail("Cover letter generation timed out (120s).")
        except FileNotFoundError:
            return self._fail(
                "Claude CLI not found. Install with: "
                "npm i -g @anthropic-ai/claude-code"
            )

        # Save to files (docx + pdf + txt)
        output_dir = Path.home() / "OneDrive" / "Documents" / "Cover Letters"
        output_dir.mkdir(parents=True, exist_ok=True)

        safe_company = re.sub(r'[^\w\s-]', '', app['company']).strip()
        safe_position = re.sub(r'[^\w\s-]', '', app['position']).strip()[:40]
        base_name = f"{safe_company} - {safe_position}"

        # Avoid overwriting
        counter = 0
        suffix = ""
        while (output_dir / f"{base_name}{suffix}.docx").exists():
            counter += 1
            suffix = f" ({counter})"

        docx_path = output_dir / f"{base_name}{suffix}.docx"
        pdf_path = output_dir / f"{base_name}{suffix}.pdf"
        txt_path = output_dir / f"{base_name}{suffix}.txt"

        # Always save plain text
        txt_path.write_text(letter, encoding="utf-8")

        # Generate DOCX
        saved_files = [str(txt_path)]
        try:
            self._write_cover_letter_docx(letter, docx_path, app)
            saved_files.append(str(docx_path))
            log.info(f"[JobTracker] DOCX saved: {docx_path.name}")

            # Convert DOCX to PDF
            try:
                import docx2pdf
                docx2pdf.convert(str(docx_path), str(pdf_path))
                saved_files.append(str(pdf_path))
                log.info(f"[JobTracker] PDF saved: {pdf_path.name}")
            except Exception as e:
                log.warning(f"[JobTracker] PDF conversion failed: {e}")
        except Exception as e:
            log.warning(f"[JobTracker] DOCX generation failed: {e}")

        # Mark cover_letter flag in DB
        self._db.update_application(app["id"], cover_letter=1)

        log.info(
            f"[JobTracker] Cover letter saved for "
            f"#{app['id']} {app['company']}"
        )

        # Build response listing saved files
        file_list = "\n".join(f"  {f}" for f in saved_files)
        return {
            "success": True,
            "response": (
                f"Cover letter for **{app['company']}** ({app['position']}) "
                f"saved to:\n\n{file_list}\n\n"
                f"DOCX and PDF are ready for upload."
            ),
            "actions_taken": [{
                "action": "cover_letter",
                "app_id": app["id"],
                "files": saved_files,
            }],
            "spoken": False,
        }

    # ── LinkedIn automation helpers ─────────────────────────────────────────

    # Lock prevents concurrent LinkedIn sessions (persistent profile
    # can only have one active Selenium driver at a time).
    _linkedin_lock = threading.Lock()
    _last_recon_ts: float = 0.0
    _RECON_COOLDOWN = 30.0  # seconds between recon calls

    @staticmethod
    def _create_linkedin_driver():
        """Create a headless Chrome driver using the persistent profile."""
        from selenium import webdriver
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.chrome.service import Service

        options = Options()
        data_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "data", "job_search_chrome_profile",
        )
        options.add_argument(f"--user-data-dir={data_dir}")
        options.add_argument("--headless=new")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument(
            "--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/146.0.0.0 Safari/537.36"
        )

        try:
            from webdriver_manager.chrome import ChromeDriverManager
            service = Service(ChromeDriverManager().install())
        except ImportError:
            service = Service()

        driver = webdriver.Chrome(service=service, options=options)
        driver.set_page_load_timeout(25)
        return driver

    @staticmethod
    def _extract_company_linkedin_url(driver, job_url: str) -> str | None:
        """Navigate to a LinkedIn job page and extract the company URL.

        Returns the base company URL (e.g. https://www.linkedin.com/company/natera)
        or None if not found.
        """
        import time as _time
        from selenium.webdriver.common.by import By

        driver.get(job_url)
        _time.sleep(4)

        # Check for login wall
        if "/login" in driver.current_url or "authwall" in driver.current_url:
            log.warning("[JobTracker] LinkedIn login required — "
                       "say 'job search login'")
            return None

        company_links = driver.find_elements(
            By.CSS_SELECTOR, 'a[href*="/company/"]')
        for link in company_links:
            href = link.get_attribute("href") or ""
            if "/company/" in href and "/jobs" not in href:
                return href.split("?")[0].rstrip("/")
        return None

    @staticmethod
    def _linkedin_im_interested(job_url: str, company: str) -> None:
        """Click 'I'm Interested' on a company's LinkedIn page.

        Runs in a background thread. Uses the persistent Chrome profile.
        """
        if not JobTrackerTalent._linkedin_lock.acquire(timeout=5):
            log.info("[JobTracker] LinkedIn session busy, skipping "
                     "'I'm Interested'")
            return
        try:
            import time as _time
            from selenium.webdriver.common.by import By

            driver = JobTrackerTalent._create_linkedin_driver()
            try:
                company_url = JobTrackerTalent._extract_company_linkedin_url(
                    driver, job_url)
                if not company_url:
                    log.info(f"[JobTracker] No company link found on "
                             f"{job_url} for 'I'm Interested'")
                    return

                life_url = f"{company_url}/life"
                driver.get(life_url)
                _time.sleep(4)

                btn = None
                for selector in (
                    'button[aria-label*="interested"]',
                    'button[aria-label*="Interested"]',
                ):
                    try:
                        btn = driver.find_element(By.CSS_SELECTOR, selector)
                        break
                    except Exception:
                        continue

                if not btn:
                    buttons = driver.find_elements(By.TAG_NAME, "button")
                    for b in buttons:
                        txt = (b.text or "").lower().strip()
                        if "interested" in txt and "not" not in txt:
                            btn = b
                            break

                if btn:
                    btn.click()
                    _time.sleep(2)
                    log.info(f"[JobTracker] Clicked 'I'm Interested' "
                             f"for {company} ({company_url})")
                else:
                    log.info(f"[JobTracker] 'I'm Interested' button not "
                             f"found for {company} ({life_url})")
            finally:
                driver.quit()
        except Exception as e:
            # Selenium's WebDriverException stringifies with the whole
            # native stack trace — usually 20+ lines of chromedriver
            # symbol addresses that are useless for a best-effort
            # background action. Trim to the first line and special-
            # case the common session-startup failure so the log stays
            # readable.
            err_text = str(e)
            short = err_text.split("\n", 1)[0].split(";", 1)[0].strip()
            if "session not created" in err_text.lower():
                log.warning(
                    f"[JobTracker] LinkedIn 'I'm Interested' skipped "
                    f"for {company} — Chrome session failed to start "
                    f"(profile conflict, version mismatch, or memory "
                    f"pressure). Status update succeeded; only the "
                    f"auto-click was skipped."
                )
            else:
                log.warning(
                    f"[JobTracker] LinkedIn 'I'm Interested' failed "
                    f"for {company}: {short}"
                )
        finally:
            JobTrackerTalent._linkedin_lock.release()

    @staticmethod
    def _discover_recruiter_and_connections(
        job_url: str, company: str, app_id: int, db_path: str,
    ) -> dict:
        """Find the recruiter and 1st-degree connections at a company.

        Single browser session, 4 phases:
        1. Extract company URL from job page
        2. Check "Meet the hiring team" on job page
        3. Navigate to company /people/ for connections
        4. Search company /people/?keywords=recruiter (if phase 2 found nothing)

        Writes results to DB. Returns a summary dict.
        """
        import json as _json
        import time as _time
        from selenium.webdriver.common.by import By

        result = {
            "recruiter_name": "",
            "recruiter_url": "",
            "connections": [],
            "error": "",
        }

        # Rate limit
        now = _time.time()
        wait = JobTrackerTalent._RECON_COOLDOWN - (
            now - JobTrackerTalent._last_recon_ts)
        if wait > 0:
            log.info(f"[JobTracker] Recon: throttling {wait:.0f}s")
            _time.sleep(wait)

        if not JobTrackerTalent._linkedin_lock.acquire(timeout=10):
            result["error"] = "Another LinkedIn task is running. Try again shortly."
            return result

        try:
            JobTrackerTalent._last_recon_ts = _time.time()
            driver = JobTrackerTalent._create_linkedin_driver()
            try:
                # ── Phase 1: extract company URL ──
                company_url = JobTrackerTalent._extract_company_linkedin_url(
                    driver, job_url)
                if not company_url:
                    result["error"] = "Could not find company LinkedIn page."
                    return result

                log.info(f"[JobTracker] Recon: company URL = {company_url}")

                # ── Phase 2: "Meet the hiring team" on job page ──
                # (We're still on the job page from phase 1)
                for selector in (
                    '.hirer-card__hirer-information a[href*="/in/"]',
                    '.jobs-poster__name a[href*="/in/"]',
                    'a[href*="/in/"][data-test-id*="hirer"]',
                ):
                    try:
                        el = driver.find_element(By.CSS_SELECTOR, selector)
                        name = el.text.strip()
                        url = (el.get_attribute("href") or "").split("?")[0]
                        if name and url:
                            result["recruiter_name"] = name
                            result["recruiter_url"] = url
                            log.info(f"[JobTracker] Recon: recruiter from "
                                     f"hiring team = {name}")
                            break
                    except Exception:
                        continue

                # Fallback: scan for "hiring team" section
                if not result["recruiter_name"]:
                    try:
                        sections = driver.find_elements(By.TAG_NAME, "section")
                        for sec in sections:
                            header = sec.text[:100].lower()
                            if "hiring" in header or "recruiter" in header:
                                links = sec.find_elements(
                                    By.CSS_SELECTOR, 'a[href*="/in/"]')
                                if links:
                                    name = links[0].text.strip()
                                    url = (links[0].get_attribute("href")
                                           or "").split("?")[0]
                                    if name:
                                        result["recruiter_name"] = name
                                        result["recruiter_url"] = url
                                        log.info(f"[JobTracker] Recon: "
                                                 f"recruiter from section "
                                                 f"= {name}")
                                break
                    except Exception:
                        pass

                # ── Phase 3: connections at company ──
                people_url = f"{company_url}/people/"
                driver.get(people_url)
                _time.sleep(4)

                connections = []
                # LinkedIn shows connection cards on the people page
                for selector in (
                    'div[data-view-name="org-people-profile-card"]',
                    '.org-people-profile-card',
                    '.artdeco-entity-lockup',
                ):
                    cards = driver.find_elements(By.CSS_SELECTOR, selector)
                    if cards:
                        break

                for card in cards[:10]:
                    try:
                        name_el = card.find_element(
                            By.CSS_SELECTOR,
                            '.artdeco-entity-lockup__title a, '
                            'a[href*="/in/"]')
                        name = name_el.text.strip()
                        url = (name_el.get_attribute("href")
                               or "").split("?")[0]
                        title = ""
                        try:
                            title_el = card.find_element(
                                By.CSS_SELECTOR,
                                '.artdeco-entity-lockup__subtitle')
                            title = title_el.text.strip()
                        except Exception:
                            pass
                        if name:
                            connections.append({
                                "name": name, "title": title, "url": url,
                            })
                    except Exception:
                        continue

                result["connections"] = connections
                if connections:
                    log.info(f"[JobTracker] Recon: {len(connections)} "
                             f"connection(s) at {company}")

                # ── Phase 4: recruiter search (if phase 2 found nothing) ──
                if not result["recruiter_name"]:
                    search_url = (f"{company_url}/people/"
                                  "?keywords=recruiter%20talent%20acquisition"
                                  "%20hiring")
                    driver.get(search_url)
                    _time.sleep(4)

                    for selector in (
                        '.artdeco-entity-lockup',
                        'div[data-view-name="org-people-profile-card"]',
                    ):
                        cards = driver.find_elements(
                            By.CSS_SELECTOR, selector)
                        if cards:
                            break

                    for card in cards[:3]:
                        try:
                            name_el = card.find_element(
                                By.CSS_SELECTOR,
                                '.artdeco-entity-lockup__title a, '
                                'a[href*="/in/"]')
                            name = name_el.text.strip()
                            url = (name_el.get_attribute("href")
                                   or "").split("?")[0]
                            title = ""
                            try:
                                title_el = card.find_element(
                                    By.CSS_SELECTOR,
                                    '.artdeco-entity-lockup__subtitle')
                                title = title_el.text.strip()
                            except Exception:
                                pass
                            if name and title:
                                title_lower = title.lower()
                                if any(kw in title_lower for kw in (
                                    "recruit", "talent", "hiring",
                                    "people", "hr ",
                                )):
                                    result["recruiter_name"] = (
                                        f"{name}, {title}")
                                    result["recruiter_url"] = url
                                    log.info(f"[JobTracker] Recon: "
                                             f"recruiter from search "
                                             f"= {name}, {title}")
                                    break
                        except Exception:
                            continue

                # ── Phase 5: write to DB ──
                try:
                    db = _DB(db_path)
                    update = {}
                    if result["recruiter_name"]:
                        update["recruiter_name"] = result["recruiter_name"]
                        update["recruiter_url"] = result["recruiter_url"]
                        update["contact_name"] = result["recruiter_name"]
                    if result["connections"]:
                        update["connections_at_co"] = _json.dumps(
                            result["connections"])
                    if update:
                        db.update_application(app_id, **update)
                        log.info(f"[JobTracker] Recon: saved results for "
                                 f"#{app_id} {company}")
                except Exception as e:
                    log.error(f"[JobTracker] Recon: DB write failed: {e}")

            finally:
                driver.quit()
        except Exception as e:
            result["error"] = str(e)
            log.warning(f"[JobTracker] Recon failed for {company}: {e}")
        finally:
            JobTrackerTalent._linkedin_lock.release()

        return result

    @staticmethod
    def _fetch_job_description(job_url: str) -> str:
        """Fetch job description text from a URL using selenium."""
        try:
            from selenium import webdriver
            from selenium.webdriver.chrome.options import Options
            from selenium.webdriver.chrome.service import Service
            from selenium.webdriver.common.by import By
            import time as _time

            options = Options()
            data_dir = os.path.join(
                os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                "data", "job_search_chrome_profile",
            )
            options.add_argument(f"--user-data-dir={data_dir}")
            options.add_argument("--headless=new")
            options.add_argument("--disable-blink-features=AutomationControlled")
            options.add_argument("--no-sandbox")
            options.add_argument("--disable-dev-shm-usage")

            try:
                from webdriver_manager.chrome import ChromeDriverManager
                service = Service(ChromeDriverManager().install())
            except ImportError:
                service = Service()

            driver = webdriver.Chrome(service=service, options=options)
            try:
                driver.get(job_url)
                _time.sleep(4)
                body = driver.find_element(By.TAG_NAME, "body").text
                # Trim to reasonable size for the prompt
                if len(body) > 5000:
                    body = body[:5000] + "\n[truncated]"
                log.info(f"[JobTracker] Fetched JD: {len(body)} chars")
                return body
            finally:
                driver.quit()
        except Exception as e:
            log.warning(f"[JobTracker] Could not fetch JD from {job_url}: {e}")
            return ""

    @staticmethod
    def _write_cover_letter_docx(
        letter_text: str, output_path: Path, app: dict
    ) -> None:
        """Write a professionally formatted cover letter DOCX."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Header: name and contact
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header.add_run("Talon User")
        run.bold = True
        run.font.size = Pt(14)
        header.paragraph_format.space_after = Pt(2)

        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = contact.add_run(
            "user@example.com | your phone | Your City, ST"
        )
        run.font.size = Pt(10)
        contact.paragraph_format.space_after = Pt(12)

        # Date
        date_para = doc.add_paragraph()
        date_para.add_run(date.today().strftime("%B %d, %Y"))
        date_para.paragraph_format.space_after = Pt(12)

        # Letter body - split into paragraphs
        paragraphs = [p.strip() for p in letter_text.split("\n\n") if p.strip()]
        for para_text in paragraphs:
            # Skip if it looks like a duplicate header/date/signature
            # that Claude may have included
            lower = para_text.lower()
            if lower.startswith("talon user"):
                continue
            if lower.startswith("dear ") or lower.startswith("re:"):
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)
                continue

            # Handle single-line breaks within a paragraph
            clean_text = para_text.replace("\n", " ")
            p = doc.add_paragraph(clean_text)
            p.paragraph_format.space_after = Pt(8)

        doc.save(str(output_path))

    # ── Follow-ups ────────────────────────────────────────────────────────────

    def _handle_follow_up_add(self, command: str, context: dict) -> dict:
        """Add a follow-up reminder for an application."""
        llm = context["llm"]
        data = self._extract_json(
            llm, _EXTRACT_FOLLOW_UP_PROMPT.format(command=command)
        )
        if not data or not data.get("company"):
            return self._fail("I need a company name. "
                              "Try: 'add a follow up for Netflix in 5 days'")

        matches = self._db.find_by_company(data["company"])
        if not matches:
            matches = self._db.search(data["company"])
        if not matches:
            return self._fail(f"No application found for '{data['company']}'.")

        app = matches[0]
        days = int(data.get("days", 3))
        due = (date.today() + timedelta(days=days)).isoformat()
        note = data.get("note", "")

        fu_id = self._db.add_follow_up(app["id"], due, note)

        return {
            "success": True,
            "response": (
                f"Follow-up #{fu_id} set for **{app['company']}** "
                f"on {_friendly_date(due)} ({due})."
                + (f"\nNote: {note}" if note else "")
            ),
            "actions_taken": [{"action": "job_follow_up_add", "id": fu_id,
                               "application_id": app["id"]}],
            "spoken": False,
        }

    def _handle_follow_up_list(self) -> dict:
        """Show pending and overdue follow-ups."""
        overdue = self._db.get_overdue_follow_ups()
        pending = self._db.get_pending_follow_ups()

        if not pending:
            return {
                "success": True,
                "response": "No pending follow-ups.",
                "actions_taken": [{"action": "job_follow_up_list"}],
                "spoken": False,
            }

        lines = [f"**Pending follow-ups** ({len(pending)}):\n"]
        for fu in pending:
            is_overdue = fu["due_date"] <= _today_iso()
            marker = " [OVERDUE]" if is_overdue else ""
            lines.append(
                f"  #{fu['id']} **{fu['company']}** -- {fu['position']} "
                f"due {_friendly_date(fu['due_date'])}{marker}"
                + (f"  ({fu['note']})" if fu["note"] else "")
            )

        return {
            "success": True,
            "response": "\n".join(lines),
            "actions_taken": [{"action": "job_follow_up_list",
                               "count": len(pending),
                               "overdue": len(overdue)}],
            "spoken": False,
        }

    # ── Export ────────────────────────────────────────────────────────────────

    def _handle_import(self, command: str, context: dict) -> dict:
        """Import applications from an existing XLSX file.

        Matches the 2026-JobSearch.xlsx format:
          Sheet 1 columns: Company, title, date applied, found on,
                           applied through, Status, Notes
        """
        try:
            import openpyxl
        except ImportError:
            return self._fail("openpyxl is not installed. Run: pip install openpyxl")

        # Try to find the file path from the command or use default
        llm = context.get("llm")
        file_path = None

        # Check for common locations
        candidates = [
            os.path.expanduser("~/OneDrive/Documents/2026-JobSearch.xlsx"),
            os.path.expanduser("~/Documents/2026-JobSearch.xlsx"),
            os.path.expanduser("~/Desktop/2026-JobSearch.xlsx"),
        ]
        for c in candidates:
            if os.path.exists(c):
                file_path = c
                break

        if not file_path:
            return self._fail(
                "Could not find 2026-JobSearch.xlsx. "
                "Checked ~/OneDrive/Documents, ~/Documents, ~/Desktop."
            )

        try:
            wb = openpyxl.load_workbook(file_path, read_only=True)
            ws = wb["Jobs Applied For"] if "Jobs Applied For" in wb.sheetnames else wb.active

            imported = 0
            skipped = 0
            for row in ws.iter_rows(min_row=2, values_only=True):
                if not row or not row[0]:
                    continue

                company = str(row[0]).strip() if row[0] else ""
                position = str(row[1]).strip() if len(row) > 1 and row[1] else ""
                if not company or not position:
                    skipped += 1
                    continue

                # Parse date
                date_applied = ""
                if len(row) > 2 and row[2]:
                    if hasattr(row[2], 'isoformat'):
                        date_applied = row[2].strftime("%Y-%m-%d")
                    else:
                        date_applied = str(row[2])[:10]

                source = str(row[3]).strip().lower() if len(row) > 3 and row[3] else ""
                method = str(row[4]).strip().lower() if len(row) > 4 and row[4] else "direct"

                # Status mapping from spreadsheet to internal
                raw_status = str(row[5]).strip().lower() if len(row) > 5 and row[5] else ""
                status = "applied"  # default
                if "reject" in raw_status:
                    status = "rejected"
                elif "interview" in raw_status or "screen" in raw_status:
                    status = "interviewing"
                elif "offer" in raw_status:
                    status = "offered"
                elif "withdraw" in raw_status:
                    status = "withdrawn"

                notes = str(row[6]).strip() if len(row) > 6 and row[6] else ""
                # The Status column often contains interview notes
                if raw_status and status == "interviewing":
                    notes = raw_status + (f"; {notes}" if notes else "")

                # Check for duplicate
                existing = self._db.search(company)
                dupe = any(
                    _normalize_company(e["company"]) == _normalize_company(company)
                    and e["position"].lower() == position.lower()
                    for e in existing
                )
                if dupe:
                    skipped += 1
                    continue

                self._db.add(
                    company=company,
                    position=position,
                    source=source,
                    method=method,
                    status=status,
                    date_applied=date_applied,
                    date_found=date_applied,
                    notes=notes,
                )
                imported += 1

            wb.close()
            return {
                "success": True,
                "response": (
                    f"Imported {imported} application(s) from {os.path.basename(file_path)}. "
                    f"{skipped} skipped (duplicates or incomplete rows)."
                ),
                "actions_taken": [{"action": "job_import", "path": file_path,
                                   "imported": imported, "skipped": skipped}],
                "spoken": False,
            }
        except Exception as e:
            return self._fail(f"Import failed: {e}")

    def _handle_export(self, context: dict) -> dict:
        """Export applications to XLSX matching the unemployment reporting format.

        Matches the user's existing 2026-JobSearch.xlsx layout:
          Sheet 1 "Jobs Applied For": Company, title, date applied, found on,
                                       applied through, Status, Notes
          Sheet 2 "Recruitment Sites Submitted": Firm, Date, Contact?
        """
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
            from openpyxl.utils import get_column_letter
        except ImportError:
            return self._fail(
                "openpyxl is not installed. Run: pip install openpyxl"
            )

        apps = self._db.list_all(include_archived=False)
        if not apps:
            return self._fail("No applications to export.")

        wb = openpyxl.Workbook()

        # ── Sheet 1: Jobs Applied For ─────────────────────────────────
        ws = wb.active
        ws.title = "Jobs Applied For"

        headers = ["Company", "title", "date applied", "found on",
                    "applied through", "Status", "Notes"]
        widths = [22, 40, 14, 14, 16, 20, 40]

        header_font = Font(bold=True)
        header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2",
                                   fill_type="solid")
        thin_border = Border(bottom=Side(style="thin"))

        for col_idx, (header, width) in enumerate(zip(headers, widths), 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center")
            ws.column_dimensions[get_column_letter(col_idx)].width = width

        for row_idx, app in enumerate(apps, 2):
            ws.cell(row=row_idx, column=1, value=app.get("company", ""))
            ws.cell(row=row_idx, column=2, value=app.get("position", ""))
            # Date as datetime for Excel formatting
            date_val = app.get("date_applied") or app.get("date_found", "")
            if date_val:
                try:
                    date_val = datetime.fromisoformat(date_val)
                except (ValueError, TypeError):
                    pass
            ws.cell(row=row_idx, column=3, value=date_val)
            ws.cell(row=row_idx, column=4, value=app.get("source", ""))
            ws.cell(row=row_idx, column=5, value=app.get("method", "direct"))
            # Status — map internal values to display
            status = app.get("status", "")
            status_display = {
                "new": "", "applied": "", "interviewing": "",
                "offered": "Offer", "rejected": "Rejection",
                "withdrawn": "Withdrawn",
            }.get(status, status)
            # Include contact/interview notes in status if present
            if status == "interviewing" and app.get("notes"):
                status_display = app["notes"]
            elif status == "rejected":
                status_display = "Rejection"
            ws.cell(row=row_idx, column=6, value=status_display)
            ws.cell(row=row_idx, column=7, value=app.get("notes", ""))

        # Freeze header row
        ws.freeze_panes = "A2"

        # ── Sheet 2: Recruitment Sites Submitted ──────────────────────
        ws2 = wb.create_sheet("Recruitment Sites Submitted")
        rec_headers = ["Firm", "Date", "Contact?"]
        rec_widths = [25, 14, 14]

        for col_idx, (header, width) in enumerate(zip(rec_headers, rec_widths), 1):
            cell = ws2.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            ws2.column_dimensions[get_column_letter(col_idx)].width = width

        # Populate from applications sourced via recruiters
        recruiter_apps = [a for a in apps
                          if (a.get("source", "").lower() in
                              ("recruiter", "staffing", "headhunter")
                              or a.get("method", "").lower() == "recruiter")]
        seen_firms = set()
        rec_row = 2
        for app in recruiter_apps:
            firm = app.get("company", "")
            if firm.lower() in seen_firms:
                continue
            seen_firms.add(firm.lower())
            ws2.cell(row=rec_row, column=1, value=firm)
            date_val = app.get("date_applied") or app.get("date_found", "")
            if date_val:
                try:
                    date_val = datetime.fromisoformat(date_val)
                except (ValueError, TypeError):
                    pass
            ws2.cell(row=rec_row, column=2, value=date_val)
            ws2.cell(row=rec_row, column=3,
                     value=app.get("contact_name", ""))
            rec_row += 1

        ws2.freeze_panes = "A2"

        # Save
        filename = f"job_tracker_export_{_today_iso()}.xlsx"
        filepath = os.path.join(_data_dir(), filename)
        wb.save(filepath)

        return {
            "success": True,
            "response": (
                f"Exported {len(apps)} application(s) to:\n`{filepath}`\n"
                f"Matches your unemployment reporting format (2 sheets)."
            ),
            "actions_taken": [{"action": "job_export", "path": filepath,
                               "count": len(apps)}],
            "spoken": False,
        }

    # ── Cowork integration ────────────────────────────────────────────────────

    def add_from_cowork(self, data: dict) -> int:
        """Add an application from Cowork bridge data.

        Args:
            data: Dict with keys matching the applications schema columns.

        Returns:
            The new application row id.
        """
        if not self._db:
            raise RuntimeError("JobTracker database not initialized")

        # Ensure required fields
        if not data.get("company") or not data.get("position"):
            raise ValueError("company and position are required")

        # Set defaults
        data.setdefault("date_found", _today_iso())
        data.setdefault("date_updated", _today_iso())
        data.setdefault("status", "new")

        # Filter to only valid columns
        valid_cols = {
            "company", "position", "location", "source", "status",
            "date_found", "date_applied", "date_updated", "contact_name",
            "contact_email", "method", "salary_range", "notes", "job_url",
            "resume_version", "cover_letter", "cowork_task_id", "fit_score",
        }
        filtered = {k: v for k, v in data.items() if k in valid_cols}
        return self._db.add_application(**filtered)

    def get_active_applications(self) -> list[dict]:
        """Return all non-archived, non-terminal applications.

        Used by Cowork to check for duplicates before adding.
        """
        if not self._db:
            return []
        return self._db.list_active()

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _extract_json(self, llm: Any, prompt: str) -> dict | None:
        """Ask the LLM to extract structured data as JSON."""
        try:
            raw = llm.generate(
                prompt,
                system_prompt=_EXTRACT_SYSTEM,
                temperature=0.1,
                max_length=256,
            )
        except LLMError as e:
            log.error(f"[JobTracker] LLM extraction failed: {e}")
            return None

        try:
            clean = raw.strip()
            if clean.startswith("```"):
                clean = re.sub(r"^```[a-z]*\n?", "", clean)
                clean = re.sub(r"\n?```$", "", clean.strip())
            match = re.search(r'\{.*\}', clean, re.DOTALL)
            if match:
                return json.loads(match.group())
        except (json.JSONDecodeError, AttributeError) as e:
            log.error(f"[JobTracker] JSON parse failed: {e}")

        return None

    # ── Claude CLI integration ─────────────────────────────────────────────

    @staticmethod
    def _claude_generate(prompt: str, timeout: float = 60.0) -> str | None:
        """Call `claude -p` for writing tasks that need better than local LLM.

        Returns the response text, or None on failure.
        Uses the Claude CLI in pipe mode — works independently of Claude Code.
        """
        claude_bin = shutil.which("claude")
        if not claude_bin:
            log.warning("[JobTracker] claude CLI not found in PATH")
            return None
        try:
            result = subprocess.run(
                [claude_bin, "-p", prompt],
                capture_output=True, text=True, timeout=timeout,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
            log.warning(f"[JobTracker] claude -p returned code {result.returncode}")
            return None
        except FileNotFoundError:
            log.warning("[JobTracker] claude CLI not found — install with: npm i -g @anthropic-ai/claude-code")
            return None
        except subprocess.TimeoutExpired:
            log.warning("[JobTracker] claude -p timed out")
            return None
        except Exception as e:
            log.error(f"[JobTracker] claude -p error: {e}")
            return None

    def evaluate_fit(self, app_id: int, job_description: str,
                     resume_summary: str) -> dict:
        """Ask Claude to evaluate how well a job matches the user's resume.

        Returns dict with fit_score (0-100) and analysis text.
        Updates the application's fit_score in the database.
        """
        prompt = (
            "You are evaluating job fit. Given a resume summary and job description, "
            "return a JSON object with:\n"
            '  {"fit_score": <0-100>, "analysis": "<2-3 sentences on fit>",'
            ' "strengths": ["..."], "gaps": ["..."]}\n\n'
            f"RESUME SUMMARY:\n{resume_summary}\n\n"
            f"JOB DESCRIPTION:\n{job_description[:3000]}\n\n"
            "Return ONLY JSON."
        )
        raw = self._claude_generate(prompt)
        if not raw:
            return {"fit_score": 0, "analysis": "Claude CLI unavailable."}

        try:
            clean = re.sub(r'^```[a-z]*\n?', '', raw.strip())
            clean = re.sub(r'\n?```$', '', clean.strip())
            m = re.search(r'\{.*\}', clean, re.DOTALL)
            if m:
                result = json.loads(m.group())
                score = int(result.get("fit_score", 0))
                # Update DB
                if self._db:
                    self._db.update(app_id, fit_score=score)
                return result
        except Exception as e:
            log.error(f"[JobTracker] Fit evaluation parse error: {e}")

        return {"fit_score": 0, "analysis": "Could not parse evaluation."}

    def draft_follow_up(self, app_id: int) -> str | None:
        """Ask Claude to draft a follow-up email for an application."""
        if not self._db:
            return None
        app = self._db.get(app_id)
        if not app:
            return None

        days_since = ""
        if app["date_applied"]:
            try:
                applied = date.fromisoformat(app["date_applied"])
                days_since = f" ({(date.today() - applied).days} days since application)"
            except ValueError:
                pass

        prompt = (
            f"Draft a brief, professional follow-up email for a job application.\n\n"
            f"Company: {app['company']}\n"
            f"Position: {app['position']}\n"
            f"Applied: {app['date_applied'] or 'recently'}{days_since}\n"
            f"Contact: {app['contact_name'] or 'Hiring Manager'}\n\n"
            "Keep it concise (3-4 sentences), professional, and enthusiastic. "
            "Do not include a subject line. Do not include a signature — the user will add their own."
        )
        return self._claude_generate(prompt)

    # ── Cowork bridge methods ────────────────────────────────────────────

    def send_to_cowork(self, task_type: str, payload: dict) -> str | None:
        """Write a task JSON for Cowork to pick up via the bridge.

        Returns the task_id, or None on failure.
        """
        import uuid

        bridge_tasks = Path.home() / "OneDrive" / "Documents" / "cowork_bridge" / "tasks"
        bridge_tasks.mkdir(parents=True, exist_ok=True)

        task_id = f"job_{task_type}_{uuid.uuid4().hex[:8]}"
        task = {
            "task_id": task_id,
            "task_type": task_type,
            "created": datetime.now().isoformat(),
            "payload": payload,
        }
        task_path = bridge_tasks / f"{task_id}.json"
        with open(task_path, "w", encoding="utf-8") as f:
            json.dump(task, f, indent=2)
        log.info(f"[JobTracker] Cowork task written: {task_path.name}")
        return task_id

    def request_resume_tailoring(self, app_id: int,
                                  job_description: str) -> str | None:
        """Send a resume tailoring request to Cowork for a specific application.

        Cowork will use the user's master resume + JD to produce tailored materials.
        Returns the Cowork task_id.
        """
        if not self._db:
            return None
        app = self._db.get(app_id)
        if not app:
            return None

        task_id = self.send_to_cowork("resume_tailor", {
            "application_id": app_id,
            "company": app["company"],
            "position": app["position"],
            "job_url": app.get("job_url", ""),
            "job_description": job_description[:5000],
            "resume_path": "~/OneDrive/Documents/resume_bullet_library.md",
            "instructions": (
                "Read the user's master resume from the resume_path. "
                "Tailor it for this position. Also draft a cover letter. "
                "Follow the writing style rules in ~/.claude/CLAUDE.md. "
                "Return both as separate sections."
            ),
        })
        if task_id and self._db:
            self._db.update(app_id, cowork_task_id=task_id)
        return task_id

    @staticmethod
    def _fail(message: str) -> dict:
        """Return a standard failure result."""
        return {
            "success": False,
            "response": message,
            "actions_taken": [],
            "spoken": False,
        }
